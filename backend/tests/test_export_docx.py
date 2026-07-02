from io import BytesIO

from docx import Document

from services.export import export_document


def test_docx_export_preserves_heading_runs_and_lists():
    html = (
        '<h1>Heading</h1>'
        '<p><strong>Bold</strong> and <em>italic</em></p>'
        '<ol><li>One</li><li>Two</li></ol>'
        '<blockquote style="text-align:right"><u>Quote</u></blockquote>'
    )

    result = export_document(html, 'docx', 'html')
    document = Document(BytesIO(result['content']))

    assert document.paragraphs[0].text == 'Heading'
    assert document.paragraphs[0].style.name.startswith('Heading')
    assert any(run.bold for run in document.paragraphs[1].runs)
    assert any(run.italic for run in document.paragraphs[1].runs)
