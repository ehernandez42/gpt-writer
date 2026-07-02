from __future__ import annotations

from dataclasses import dataclass, field
from html import escape
from html.parser import HTMLParser
from io import BytesIO
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_CONTENT_TYPE = "application/pdf"
ALIGNMENTS = {"left", "center", "right", "justify"}
HEADING_TAGS = {"h1": 1, "h2": 2}
PDF_ALIGNMENT_MAP = {
    "left": TA_LEFT,
    "center": TA_CENTER,
    "right": TA_RIGHT,
    "justify": TA_JUSTIFY,
}
DOCX_ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


@dataclass
class TextRun:
    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False


@dataclass
class Block:
    type: str
    inlines: list[TextRun] = field(default_factory=list)
    level: int | None = None
    align: str = "left"
    ordered: bool | None = None
    items: list[list[TextRun]] = field(default_factory=list)


class ExportHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[Block] = []
        self.inline_state = {"bold": False, "italic": False, "underline": False}
        self.block_stack: list[Block] = []
        self.list_stack: list[Block] = []
        self.current_list_item: list[TextRun] | None = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attrs_dict = dict(attrs)

        if tag in {"strong", "b"}:
            self.inline_state["bold"] = True
            return
        if tag in {"em", "i"}:
            self.inline_state["italic"] = True
            return
        if tag == "u":
            self.inline_state["underline"] = True
            return
        if tag == "br":
            self._append_text("\n")
            return
        if tag == "hr":
            self.blocks.append(Block(type="horizontal_rule"))
            return
        if tag in HEADING_TAGS:
            self.block_stack.append(
                Block(
                    type="heading",
                    level=HEADING_TAGS[tag],
                    align=_parse_alignment(attrs_dict.get("style")),
                )
            )
            return
        if tag == "p":
            self.block_stack.append(Block(type="paragraph", align=_parse_alignment(attrs_dict.get("style"))))
            return
        if tag == "blockquote":
            self.block_stack.append(Block(type="blockquote", align=_parse_alignment(attrs_dict.get("style"))))
            return
        if tag == "ul":
            self.list_stack.append(Block(type="list", ordered=False))
            return
        if tag == "ol":
            self.list_stack.append(Block(type="list", ordered=True))
            return
        if tag == "li":
            self.current_list_item = []

    def handle_endtag(self, tag: str) -> None:
        if tag in {"strong", "b"}:
            self.inline_state["bold"] = False
            return
        if tag in {"em", "i"}:
            self.inline_state["italic"] = False
            return
        if tag == "u":
            self.inline_state["underline"] = False
            return
        if tag in HEADING_TAGS or tag in {"p", "blockquote"}:
            if self.block_stack:
                block = self.block_stack.pop()
                if block.inlines:
                    self.blocks.append(block)
            return
        if tag == "li":
            if self.list_stack and self.current_list_item is not None:
                cleaned = _normalize_runs(self.current_list_item)
                if cleaned:
                    self.list_stack[-1].items.append(cleaned)
            self.current_list_item = None
            return
        if tag in {"ul", "ol"}:
            if self.list_stack:
                block = self.list_stack.pop()
                if block.items:
                    self.blocks.append(block)

    def handle_data(self, data: str) -> None:
        self._append_text(data)

    def _append_text(self, text: str) -> None:
        if not text:
            return

        if self.current_list_item is not None:
            target = self.current_list_item
        elif self.block_stack:
            target = self.block_stack[-1].inlines
        else:
            if not text.strip():
                return
            block = Block(type="paragraph")
            self.blocks.append(block)
            self.block_stack.append(block)
            target = block.inlines

        run = TextRun(text=text, **self.inline_state)
        if target and _same_style(target[-1], run):
            target[-1].text += run.text
        else:
            target.append(run)


def _same_style(left: TextRun, right: TextRun) -> bool:
    return left.bold == right.bold and left.italic == right.italic and left.underline == right.underline


def _parse_alignment(style: str | None) -> str:
    if not style:
        return "left"
    match = re.search(r"text-align\s*:\s*(left|center|right|justify)", style, re.IGNORECASE)
    if not match:
        return "left"
    value = match.group(1).lower()
    return value if value in ALIGNMENTS else "left"


def _normalize_runs(runs: list[TextRun]) -> list[TextRun]:
    normalized: list[TextRun] = []
    for run in runs:
        if not run.text:
            continue
        if normalized and _same_style(normalized[-1], run):
            normalized[-1].text += run.text
        else:
            normalized.append(TextRun(run.text, run.bold, run.italic, run.underline))
    return normalized


def parse_html_document(html: str) -> list[Block]:
    parser = ExportHTMLParser()
    parser.feed(html)
    parser.close()

    blocks: list[Block] = []
    seen_ids: set[int] = set()
    for block in parser.blocks:
        if id(block) in seen_ids:
            continue
        seen_ids.add(id(block))
        if block.type == "list":
            if block.items:
                blocks.append(block)
            continue
        block.inlines = _normalize_runs(block.inlines)
        if block.type == "horizontal_rule" or block.inlines:
            blocks.append(block)
    return blocks


def _runs_to_plain_text(runs: list[TextRun]) -> str:
    return "".join(run.text for run in runs)


def _runs_to_reportlab_markup(runs: list[TextRun]) -> str:
    parts: list[str] = []
    for run in runs:
        text = escape(run.text).replace("\n", "<br/>")
        if run.bold:
            text = f"<b>{text}</b>"
        if run.italic:
            text = f"<i>{text}</i>"
        if run.underline:
            text = f"<u>{text}</u>"
        parts.append(text)
    return "".join(parts) or " "


def _pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    body = ParagraphStyle(
        "ExportBody",
        parent=base["BodyText"],
        fontName="Helvetica",
        fontSize=11,
        leading=14,
        spaceAfter=0,
    )
    return {
        "paragraph": body,
        "heading1": ParagraphStyle("ExportHeading1", parent=base["Heading1"], spaceAfter=0),
        "heading2": ParagraphStyle("ExportHeading2", parent=base["Heading2"], spaceAfter=0),
        "blockquote": ParagraphStyle(
            "ExportBlockquote",
            parent=body,
            leftIndent=18,
            textColor="#444444",
        ),
    }


def _style_for_block(block: Block, styles: dict[str, ParagraphStyle]) -> ParagraphStyle:
    if block.type == "heading":
        base_style = styles[f"heading{block.level or 1}"]
    elif block.type == "blockquote":
        base_style = styles["blockquote"]
    else:
        base_style = styles["paragraph"]
    return ParagraphStyle(
        f"{base_style.name}-{block.align}",
        parent=base_style,
        alignment=PDF_ALIGNMENT_MAP[block.align],
    )


def export_pdf_document(document: list[Block]) -> dict:
    buffer = BytesIO()
    template = SimpleDocTemplate(
        buffer,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = _pdf_styles()
    flowables = []

    for block in document:
        if block.type == "horizontal_rule":
            flowables.append(HRFlowable(width="100%"))
            flowables.append(Spacer(1, 12))
            continue

        if block.type == "list":
            list_items = []
            for item in block.items:
                paragraph = Paragraph(_runs_to_reportlab_markup(item), styles["paragraph"])
                list_items.append(ListItem(paragraph))
            flowables.append(
                ListFlowable(
                    list_items,
                    bulletType="1" if block.ordered else "bullet",
                    leftIndent=18,
                )
            )
            flowables.append(Spacer(1, 12))
            continue

        flowables.append(Paragraph(_runs_to_reportlab_markup(block.inlines), _style_for_block(block, styles)))
        flowables.append(Spacer(1, 12))

    template.build(flowables)
    return {
        "filename": "generated.pdf",
        "content_type": PDF_CONTENT_TYPE,
        "content": buffer.getvalue(),
    }


def _apply_docx_alignment(paragraph, align: str) -> None:
    paragraph.alignment = DOCX_ALIGNMENT_MAP[align]


def _add_docx_runs(paragraph, runs: list[TextRun]) -> None:
    for text_run in runs:
        run = paragraph.add_run(text_run.text)
        run.bold = text_run.bold
        run.italic = text_run.italic
        run.underline = text_run.underline


def export_docx_document(document: list[Block]) -> dict:
    doc = Document()

    for block in document:
        if block.type == "heading":
            paragraph = doc.add_paragraph(style=f"Heading {block.level or 1}")
            _add_docx_runs(paragraph, block.inlines)
            _apply_docx_alignment(paragraph, block.align)
            continue

        if block.type == "blockquote":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.5)
            _add_docx_runs(paragraph, block.inlines)
            _apply_docx_alignment(paragraph, block.align)
            continue

        if block.type == "list":
            style_name = "List Number" if block.ordered else "List Bullet"
            for item in block.items:
                paragraph = doc.add_paragraph(style=style_name)
                _add_docx_runs(paragraph, item)
            continue

        if block.type == "horizontal_rule":
            doc.add_paragraph("―" * 16)
            continue

        paragraph = doc.add_paragraph()
        _add_docx_runs(paragraph, block.inlines)
        _apply_docx_alignment(paragraph, block.align)

    buffer = BytesIO()
    doc.save(buffer)
    return {
        "filename": "generated.docx",
        "content_type": DOCX_CONTENT_TYPE,
        "content": buffer.getvalue(),
    }


def export_document(content: str, format: str, content_type: str) -> dict:
    if content_type != "html":
        raise ValueError("Unsupported content type")

    document = parse_html_document(content)
    if format == "pdf":
        return export_pdf_document(document)
    if format == "docx":
        return export_docx_document(document)
    raise ValueError("Unsupported export format")


def export_text(text: str, format: str) -> dict:
    if format == "docx":
        buffer = BytesIO()
        document = Document()
        document.add_paragraph(text)
        document.save(buffer)
        return {
            "filename": "generated.docx",
            "content_type": DOCX_CONTENT_TYPE,
            "content": buffer.getvalue(),
        }

    if format == "pdf":
        return export_pdf_document([Block(type="paragraph", inlines=[TextRun(text=text)])])

    raise ValueError("Unsupported export format")
